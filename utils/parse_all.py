"""通用解析脚本：将 raw 文件解析为 Markdown + 元数据 JSON.

用法：
    python utils/parse_all.py                           # 解析所有文件
    python utils/parse_all.py --source sse               # 仅上交所
    python utils/parse_all.py --source sse --limit 5     # 上交所前5个
    python utils/parse_all.py --only sse/技术通知        # 仅指定路径
"""

from __future__ import annotations

import argparse
import hashlib
import json
import logging
import re
import sys
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

logging.basicConfig(
    level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
    handlers=[
        logging.FileHandler(Path("log/parse_all.log"), encoding="utf-8"),
        logging.StreamHandler(),
    ],
)
logger = logging.getLogger("parse_all")

RAW_DIR = Path("knowledge/raw")
OUT_DIR = Path("knowledge/articles")

DOC_TYPES = {
    "规格说明书": "interface_spec",
    "接口规范": "interface_spec",
    "数据接口": "interface_spec",
    "接口规格": "interface_spec",
    "技术实施指南": "guide",
    "实施指南": "guide",
    "技术指南": "guide",
    "通知": "technical_notice",
    "测试方案": "test_doc",
    "测试": "test_doc",
    "软件": "software",
    "下载": "software",
    "错误代码表": "software",
    "杂志": "magazine",
    "业务规则": "business_rule",
}


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _infer_doc_type(title: str, category: str) -> str:
    for kw, dt in DOC_TYPES.items():
        if kw in title:
            return dt
    cat_map = {
        "技术通知": "technical_notice", "服务指引": "guide",
        "技术接口": "interface_spec", "技术指南": "guide",
        "软件下载": "software", "测试文档": "test_doc",
        "技术杂志": "magazine", "历史资料": "guide",
        "技术公告": "technical_notice", "交易系统介绍": "guide",
        "数据接口": "interface_spec", "业务规则": "business_rule",
    }
    return cat_map.get(category, "technical_notice")


def _infer_version(title: str) -> str | None:
    patterns = [
        r"V(\d+[\.\d]*)", r"v(\d+[\.\d]*)", r"Ver(\d+[\.\d]*)",
        r"ver(\d+[\.\d]*)", r"(\d+[\.\d]*)版", r"版本(\d+[\.\d]*)",
    ]
    for p in patterns:
        m = re.search(p, title)
        if m:
            return m.group(1)
    return None


def _extract_date_from_text(text: str) -> str | None:
    patterns = [
        r"(\d{4})年(\d{1,2})月(\d{1,2})日",
        r"(\d{4})-(\d{2})-(\d{2})",
        r"(\d{4})\.(\d{2})\.(\d{2})",
    ]
    for p in patterns:
        m = re.search(p, text)
        if m:
            return f"{m.group(1)}-{m.group(2).zfill(2)}-{m.group(3).zfill(2)}"
    return None


def _color_to_span(color: int) -> str:
    c = color & 0xFFFFFF
    r, g, b = (c >> 16) & 0xFF, (c >> 8) & 0xFF, c & 0xFF
    if r > 200 and g < 80 and b < 80:
        return "red"
    if b > 200 and r < 80 and g < 80:
        return "blue"
    if r < 60 and g < 60 and b < 60:
        return None
    return None


def _detect_change_annotation(text: str) -> str | None:
    stripped = text.strip()
    for prefix in ["【新增】", "【修改】", "【删除】", "【废止】", "新增", "修改", "删除", "废止"]:
        if stripped.startswith(prefix) or stripped.startswith(f"［{prefix}］"):
            return prefix.strip("【】")
    return None


def parse_pdf(path: Path, source: str, category: str) -> dict[str, Any]:
    import fitz
    title = Path(path.stem).stem
    author = ""
    full_text_parts: list[str] = []
    changes_found = False
    page_count = 0
    doc = None
    try:
        doc = fitz.open(path)
        try:
            meta = doc.metadata
            if meta.get("title"):
                title = meta["title"]
            author = meta.get("author", "")
        except Exception:
            pass
        page_count = doc.page_count
        for page_num, page in enumerate(doc):
            text_blocks = page.get_text("dict", sort=True)
            page_parts: list[str] = []
            for block in text_blocks.get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    line_text = ""
                    line_spans: list[str] = []
                    for span in line.get("spans", []):
                        span_text = span.get("text", "").strip()
                        if not span_text:
                            continue
                        color = span.get("color", 0)
                        css_color = _color_to_span(color)
                        annotation = _detect_change_annotation(span_text)
                        if css_color or annotation:
                            changes_found = True
                        if css_color:
                            prefix = f"[{annotation}] " if annotation else ""
                            line_spans.append(f'<span style="color:{css_color}">{prefix}{span_text}</span>')
                        else:
                            line_spans.append(span_text)
                        line_text = " ".join(line_spans)
                    if line_text:
                        page_parts.append(line_text)
            page_text = "\n".join(page_parts)
            if page_text.strip():
                full_text_parts.append(page_text)
    except Exception as e:
        logger.error("PDF parse error [%s]: %s", path.name, e)
        raise
    finally:
        if doc:
            doc.close()
    markdown = "\n\n".join(full_text_parts)
    if not markdown.strip():
        markdown = f"*[无法从 PDF 提取文本 — 可能是扫描件: {path.name}]*"
    return {
        "markdown": markdown,
        "title": title,
        "author": author,
        "changes_found": changes_found,
        "page_count": page_count,
    }


def parse_docx(path: Path, source: str, category: str) -> dict[str, Any]:
    import docx as docx_lib
    d = docx_lib.Document(path)
    title = Path(path.stem).stem
    full_text: list[str] = []
    changes_found = False
    for para in d.paragraphs:
        style = para.style.name if para.style else "Normal"
        text = para.text.strip()
        if not text:
            continue
        prefix = "#" if "Heading 1" in style else "##" if "Heading" in style else ""
        if prefix:
            full_text.append(f"{prefix} {text}")
        else:
            annotation = _detect_change_annotation(text)
            if annotation:
                changes_found = True
                full_text.append(f'<span style="color:red">[{annotation}] {text}</span>')
            else:
                full_text.append(text)
    if d.tables:
        for table in d.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
            if rows:
                header = rows[0]
                sep = "| " + " | ".join(["---"] * len(table.columns)) + " |"
                full_text.append("\n" + header + "\n" + sep + "\n" + "\n".join(rows[1:]))
    return {
        "markdown": "\n\n".join(full_text),
        "title": title,
        "changes_found": changes_found,
    }


def parse_xlsx(path: Path, source: str, category: str) -> dict[str, Any]:
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    title = Path(path.stem).stem
    full_parts: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        full_parts.append(f"## Sheet: {sheet_name}")
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        header = rows[0]
        sep = "| " + " | ".join(["---"] * len(header)) + " |"
        full_parts.append("| " + " | ".join(str(c or "") for c in header) + " |")
        full_parts.append(sep)
        for row in rows[1:]:
            full_parts.append("| " + " | ".join(str(c or "") for c in row) + " |")
    wb.close()
    return {
        "markdown": "\n".join(full_parts),
        "title": title,
        "changes_found": False,
    }


def parse_html_text(html: str) -> str:
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, "lxml")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()
    parts: list[str] = []
    for tag in soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "table"]):
        text = tag.get_text(strip=True)
        if not text:
            continue
        if tag.name in ("h1", "h2", "h3", "h4"):
            prefix = "#" * int(tag.name[1])
            parts.append(f"{prefix} {text}")
        elif tag.name == "li":
            parts.append(f"- {text}")
        else:
            parts.append(text)
    return "\n\n".join(parts)


def build_meta(
    raw_path: Path,
    md_path: Path,
    parse_result: dict[str, Any],
    source: str,
    category: str,
    sub_category: str | None,
    crawl_item: dict[str, Any] | None,
) -> dict[str, Any]:
    file_hash = _sha256(raw_path)
    title = parse_result.get("title", raw_path.stem)
    doc_type = _infer_doc_type(title, category)
    version = _infer_version(title)
    public_date = crawl_item.get("publish_date") if crawl_item else None
    if not public_date:
        text = parse_result.get("markdown", "")[:2000]
        public_date = _extract_date_from_text(text)
    return {
        "title": title,
        "source_url": crawl_item.get("url") if crawl_item else None,
        "raw_path": str(raw_path),
        "markdown_path": str(md_path),
        "file_hash": f"sha256:{file_hash}",
        "file_format": raw_path.suffix.lstrip("."),
        "page_count": parse_result.get("page_count", 0),
        "doc_type": doc_type,
        "version": version,
        "previous_version": None,
        "public_date": public_date,
        "effective_date": None,
        "has_changes": parse_result.get("changes_found", False),
        "parse_status": "success",
        "parse_date": datetime.now(timezone.utc).isoformat(),
    }


def collect_raw_files(source_filter: str | None = None, only_path: str | None = None, limit: int | None = None) -> list[tuple[str, str, str | None, Path]]:
    files: list[tuple[str, str, str | None, Path]] = []
    known_exts = {".pdf", ".docx", ".xlsx", ".xls", ".zip", ".html", ".shtml"}
    for raw_dir in RAW_DIR.iterdir():
        if not raw_dir.is_dir():
            continue
        source = raw_dir.name
        if source_filter and source != source_filter:
            continue
        for cat_dir in raw_dir.iterdir():
            if not cat_dir.is_dir():
                continue
            category = cat_dir.name
            sub_category = None
            if source == "chinaclear":
                sub_category = category
                category = "业务规则"
            for f in sorted(cat_dir.iterdir()):
                if f.is_file() and f.suffix.lower() in known_exts:
                    if only_path and only_path not in str(f):
                        continue
                    files.append((source, category, sub_category, f))
    if limit:
        files = files[:limit]
    return files


def parse_file(source: str, category: str, sub_category: str | None, path: Path) -> dict[str, Any] | None:
    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            return parse_pdf(path, source, category)
        elif ext in (".docx",):
            return parse_docx(path, source, category)
        elif ext in (".xlsx", ".xls"):
            return parse_xlsx(path, source, category)
        elif ext in (".html", ".shtml"):
            text = path.read_text(encoding="utf-8", errors="replace")
            md = parse_html_text(text)
            return {"markdown": md, "title": path.stem, "changes_found": False}
        elif ext == ".zip":
            return {"markdown": f"*[ZIP archive: {path.name} — needs recursive extraction]*", "title": path.stem, "changes_found": False}
        else:
            logger.warning("Unsupported format: %s", path)
            return None
    except Exception as e:
        logger.error("Parse failed [%s]: %s", path.name, e)
        return None


def save_output(source: str, category: str, sub_category: str | None, path: Path, result: dict[str, Any], crawl_meta: dict[str, Any] | None):
    date_part = re.sub(r"^\d{8}_", "", path.stem)[:60]
    safe_name = re.sub(r'[<>:"/\\|?*]', "_", date_part)

    rel_cat = sub_category if sub_category else category
    md_dir = OUT_DIR / source / "markdown" / rel_cat
    meta_dir = OUT_DIR / source / "metadata" / rel_cat
    md_dir.mkdir(parents=True, exist_ok=True)
    meta_dir.mkdir(parents=True, exist_ok=True)

    md_path = md_dir / f"{safe_name}.md"
    meta_path = meta_dir / f"{safe_name}_meta.json"

    meta = build_meta(path, md_path, result, source, category, sub_category, crawl_meta)
    meta["sub_category"] = sub_category

    md_content = result["markdown"]
    if meta["has_changes"]:
        md_content += f'\n\n> **变更标注说明**：本文档中已用 `<span style="color:...">` 标注了变更内容（红色=修改/新增，蓝色=其他说明）。\n'
    md_content += f'\n\n<metadata>\n{json.dumps(meta, ensure_ascii=False, indent=2)}\n</metadata>'

    md_path.write_text(md_content, encoding="utf-8")
    meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")
    return md_path, meta_path


def load_crawl_meta() -> list[dict[str, Any]]:
    path = RAW_DIR / "crawl_metadata.json"
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return []


def match_crawl_meta(path: Path, crawl_items: list[dict[str, Any]]) -> dict[str, Any] | None:
    fname = path.name
    for item in crawl_items:
        if item.get("local_path") and fname in item["local_path"]:
            return item
    return None


def main():
    parser = argparse.ArgumentParser(description="解析所有原始文件为 Markdown + 元数据")
    parser.add_argument("--source", help="限定数据源: sse / szse / chinaclear")
    parser.add_argument("--only", help="仅解析包含此路径的文件")
    parser.add_argument("--limit", type=int, help="最大解析文件数")
    parser.add_argument("--request-delay", type=float, default=0.3, help="文件间延迟(秒)")
    args = parser.parse_args()

    logger.info("Scanning raw files...")
    files = collect_raw_files(args.source, args.only, args.limit)
    logger.info("Found %d raw files to parse", len(files))

    crawl_items = load_crawl_meta()

    ok = fail = 0
    for source, category, sub_category, fpath in files:
        logger.info("Parsing [%s] [%s] %s", source, category, fpath.name)
        crawl_item = match_crawl_meta(fpath, crawl_items)
        result = parse_file(source, category, sub_category, fpath)
        if result is None:
            fail += 1
            continue
        try:
            md_path, meta_path = save_output(source, category, sub_category, fpath, result, crawl_item)
            logger.info("  -> %s", md_path)
            ok += 1
        except Exception as e:
            logger.error("Save failed [%s]: %s", fpath.name, e)
            fail += 1
        time.sleep(args.request_delay)

    logger.info("Done: %d success, %d failed", ok, fail)


if __name__ == "__main__":
    main()
